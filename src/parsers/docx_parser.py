"""DOCX 문서 파서"""

from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.table import Table

from .base_parser import BaseParser
from ..utils.logger import get_logger

logger = get_logger("docx_parser")


class DOCXParser(BaseParser):
    """DOCX 문서 파서"""

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]

    def parse(self, file_path: Path) -> Dict[str, Any]:
        """
        DOCX를 파싱하여 구조화된 데이터 반환

        Args:
            file_path: DOCX 파일 경로

        Returns:
            파싱된 데이터 딕셔너리
        """
        logger.info(f"DOCX 파싱 시작: {file_path}")

        doc = Document(file_path)

        result = {
            "raw_text": self.extract_text(file_path),
            "tables": self.extract_tables(file_path),
            "sections": self._extract_sections(doc),
            "metadata": self._extract_metadata(doc),
            "styles": self._extract_styles(doc),
        }

        logger.info(
            f"DOCX 파싱 완료: {len(result['raw_text'])} 문자, "
            f"{len(result['tables'])} 테이블, "
            f"{len(result['sections'])} 섹션"
        )

        return result

    def extract_text(self, file_path: Path) -> str:
        """전체 텍스트 추출"""
        try:
            doc = Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error(f"텍스트 추출 실패: {e}")
            return ""

    def extract_tables(self, file_path: Path) -> List[Dict[str, Any]]:
        """테이블 데이터 추출"""
        tables = []

        try:
            doc = Document(file_path)

            for i, table in enumerate(doc.tables):
                table_data = self._table_to_dict(table, i)
                if table_data:
                    tables.append(table_data)
        except Exception as e:
            logger.error(f"테이블 추출 실패: {e}")

        return tables

    def _table_to_dict(self, table: Table, index: int) -> Dict[str, Any]:
        """Table 객체를 딕셔너리로 변환"""
        rows = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            rows.append(row_data)

        if not rows:
            return {}

        # 첫 번째 행을 헤더로 처리
        headers = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []

        return {
            "table_index": index,
            "headers": headers,
            "rows": data_rows,
            "raw_data": rows,
        }

    def _extract_sections(self, doc: Document) -> List[Dict[str, Any]]:
        """헤딩 기반 섹션 추출"""
        sections = []
        current_section = {"title": "", "content": [], "level": 0}

        for para in doc.paragraphs:
            # 헤딩 스타일 체크
            if para.style and para.style.name.startswith("Heading"):
                # 이전 섹션 저장
                if current_section["content"] or current_section["title"]:
                    sections.append(current_section)

                # 헤딩 레벨 추출
                level = 1
                if para.style.name[-1].isdigit():
                    level = int(para.style.name[-1])

                current_section = {
                    "title": para.text.strip(),
                    "content": [],
                    "level": level,
                    "style": para.style.name,
                }
            else:
                text = para.text.strip()
                if text:
                    current_section["content"].append(text)

        # 마지막 섹션 저장
        if current_section["content"] or current_section["title"]:
            sections.append(current_section)

        return sections

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """메타데이터 추출"""
        try:
            core_props = doc.core_properties
            return {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
        except Exception as e:
            logger.warning(f"메타데이터 추출 실패: {e}")
            return {}

    def _extract_styles(self, doc: Document) -> Dict[str, Any]:
        """사용된 스타일 추출"""
        styles_used = set()

        for para in doc.paragraphs:
            if para.style:
                styles_used.add(para.style.name)

        return {"styles_used": list(styles_used)}
